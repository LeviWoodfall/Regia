"""
Document processing handlers for non-PDF file types.
Supports: DOCX, XLSX, images (JPG, PNG, TIFF, BMP, GIF, WEBP).
"""

import io
import logging
from typing import Optional

logger = logging.getLogger("regia.processing.doc_handler")


# === DOCX Handler ===

def extract_docx_text(filepath: str) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed — cannot extract DOCX text")
        return ""
    except Exception as e:
        logger.error(f"Failed to extract DOCX text from {filepath}: {e}")
        return ""


def get_docx_metadata(filepath: str) -> dict:
    """Extract metadata from a DOCX file."""
    try:
        from docx import Document
        doc = Document(filepath)
        props = doc.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "subject": props.subject or "",
        }
    except Exception:
        return {}


def get_docx_page_count(filepath: str) -> int:
    """Estimate page count for a DOCX file (based on paragraph count)."""
    try:
        from docx import Document
        doc = Document(filepath)
        # Rough estimate: ~25 paragraphs per page
        para_count = len([p for p in doc.paragraphs if p.text.strip()])
        return max(1, para_count // 25 + 1)
    except Exception:
        return 1


# === XLSX Handler ===

def extract_xlsx_text(filepath: str) -> str:
    """Extract text from an XLSX file (all sheets)."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    parts.append(row_text)
        wb.close()
        return "\n".join(parts)
    except ImportError:
        logger.warning("openpyxl not installed — cannot extract XLSX text")
        return ""
    except Exception as e:
        logger.error(f"Failed to extract XLSX text from {filepath}: {e}")
        return ""


def get_xlsx_metadata(filepath: str) -> dict:
    """Extract metadata from an XLSX file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, read_only=True)
        props = wb.properties
        meta = {
            "title": props.title or "",
            "creator": props.creator or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "sheet_count": len(wb.sheetnames),
            "sheet_names": wb.sheetnames,
        }
        wb.close()
        return meta
    except Exception:
        return {}


def get_xlsx_sheet_count(filepath: str) -> int:
    """Get the number of sheets in an XLSX file."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, read_only=True)
        count = len(wb.sheetnames)
        wb.close()
        return count
    except Exception:
        return 1


# === Image Handler ===

def extract_image_text(filepath: str, ocr_config=None) -> str:
    """Extract text from an image via OCR."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("pytesseract/Pillow not installed — cannot OCR images")
        return ""

    try:
        image = Image.open(filepath)
        lang = ocr_config.language if ocr_config else "eng"
        text = pytesseract.image_to_string(image, lang=lang)
        return text
    except Exception as e:
        logger.error(f"Image OCR failed for {filepath}: {e}")
        return ""


def get_image_metadata(filepath: str) -> dict:
    """Extract metadata from an image file."""
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS
        img = Image.open(filepath)
        meta = {
            "width": img.width,
            "height": img.height,
            "format": img.format or "",
            "mode": img.mode,
        }
        # Extract EXIF if available
        exif = img.getexif()
        if exif:
            for tag_id, value in exif.items():
                tag = TAGS.get(tag_id, tag_id)
                if isinstance(value, (str, int, float)):
                    meta[f"exif_{tag}"] = value
        img.close()
        return meta
    except Exception:
        return {}


def render_image_preview(filepath: str, max_size: int = 800) -> Optional[bytes]:
    """Render a thumbnail preview of an image as PNG bytes."""
    try:
        from PIL import Image
        img = Image.open(filepath)
        img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception as e:
        logger.error(f"Failed to render image preview for {filepath}: {e}")
        return None


# === File Type Detection ===

# Supported MIME types and extensions
SUPPORTED_TYPES = {
    # PDF
    "application/pdf": "pdf",
    # DOCX
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    # XLSX
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    # Images
    "image/jpeg": "image",
    "image/png": "image",
    "image/tiff": "image",
    "image/bmp": "image",
    "image/gif": "image",
    "image/webp": "image",
}

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
    ".tiff": "image",
    ".tif": "image",
    ".bmp": "image",
    ".gif": "image",
    ".webp": "image",
}


def detect_file_type(filename: str, content_type: str = "") -> Optional[str]:
    """Detect file type from filename or MIME type. Returns: 'pdf', 'docx', 'xlsx', 'image', or None."""
    # Check MIME type first
    if content_type and content_type in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[content_type]
    # Check extension
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return SUPPORTED_EXTENSIONS.get(ext)
